"""
Convert a data frame to word.
"""
import pandas as pd
from docx import Document


def frame_to_word(frame: pd.DataFrame, file):
    """
    Write a pandas.DataFrame to a word document.
    Will have the structure:

    index[0] [h2]
        column[0] [h3]
        ...
        column[1] [h3]
        ...
        column[2] [h3]
        ...
        etc
    index[1] [h2]
        column[0] [h3]
        ...
    etc.

    Parameters
    ----------
    frame: pandas.DataFrame
        The data set to write to file
    file: str
        The file path to write the .doc file to
    Returns
    -------
    None
    """

    doc = Document()
    _iter = frame.fillna(" ").iterrows()
    for index, row in _iter:
        doc.add_heading(index, 2)
        for col, content in row.items():
            doc.add_heading(col, 3)
            for item in content.split("\n"):
                doc.add_paragraph(item)

    doc.save(file)
